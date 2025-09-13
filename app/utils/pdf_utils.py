from typing import Optional
from collections import Counter
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from io import BytesIO
from docx import Document
from app.core.logger import get_logger
import re

logger = get_logger(__name__)

def extract_text_from_pdf(file_bytes: bytes) -> Optional[str]:
    """
    Extracts text from PDF using PyMuPDF.
    Returns None if extraction fails or text is empty.
    """
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            text_chunks = [page.get_text("text").strip() for page in pdf]
        extracted_text = "\n".join([t for t in text_chunks if t])

        return extracted_text if extracted_text else None
    except Exception:
        return None

def extract_text_from_docx(file_bytes: bytes) -> Optional[str]:
    """
    Extracts text from DOCX resumes.
    Includes paragraphs and tables.
    Returns None if extraction fails or file is empty.
    """
    try:
        doc = Document(BytesIO(file_bytes))
        text_chunks = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_chunks.append(para.text.strip())

        # Extract tables (cell by cell)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_chunks.append(cell.text.strip())

        extracted_text = "\n".join(text_chunks)
        return extracted_text if extracted_text else None
    except Exception:
        return None


def extract_resume_text(file_bytes: bytes, content_type: str, filename: str) -> str:
    """
    Universal extractor for resumes (PDF or DOCX).
    """
    try:
        fname = filename.lower()
        ctype = (content_type or "").lower()

        if "pdf" in ctype or fname.endswith(".pdf"):
            return extract_text_from_pdf(file_bytes)
        elif "word" in ctype or fname.endswith(".docx"):
            return extract_text_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {filename}")
    except Exception as e:
        logger.error(f"Resume extraction failed: {e}")
        return ""


def extract_keywords(job_description: str, top_n: int = 30) -> list:
    """
    Extracts keywords (skills, tools, terms) from JD deterministically.
    Very simple version using regex + frequency filtering.
    """
    # Normalize
    text = job_description.lower()
    
    # Keep only words
    words = re.findall(r"[a-zA-Z]+", text)
    
    # Common stopwords (you can extend this list)
    stopwords = {"the","and","with","for","a","an","to","of","in","on","at","is","as","by","or","be","from"}
    
    # Filter
    filtered = [w for w in words if w not in stopwords and len(w) > 2]
    
    # Count frequency
    counts = Counter(filtered)
    
    # Pick top-N terms
    keywords = [word for word, _ in counts.most_common(top_n)]
    
    return keywords